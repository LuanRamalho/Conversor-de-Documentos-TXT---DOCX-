import os
from tkinter import Tk, filedialog, StringVar, Label, Button, OptionMenu, messagebox
from docx import Document


def select_file():
    """Permite ao usuário selecionar um arquivo para conversão."""
    file_path = filedialog.askopenfilename(
        title="Selecione o arquivo",
        filetypes=(("Arquivos de Texto", "*.txt"), ("Arquivos DOCX", "*.docx")),
    )
    if file_path:
        input_file.set(file_path)
        lbl_selected_file["text"] = f"Arquivo selecionado: {os.path.basename(file_path)}"


def convert_file():
    """Converte o arquivo de entrada para o formato desejado."""
    file_path = input_file.get()
    input_ext = input_format.get()
    output_ext = output_format.get()

    if not file_path:
        messagebox.showerror("Erro", "Nenhum arquivo foi selecionado.")
        return

    if input_ext == output_ext:
        messagebox.showerror("Erro", "Os formatos de entrada e saída são iguais.")
        return

    try:
        if input_ext == "TXT" and output_ext == "DOCX":
            with open(file_path, "r", encoding="utf-8") as txt_file:
                content = txt_file.read()
            doc = Document()
            doc.add_paragraph(content)
            output_path = file_path.rsplit(".", 1)[0] + ".docx"
            doc.save(output_path)

        elif input_ext == "DOCX" and output_ext == "TXT":
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            output_path = file_path.rsplit(".", 1)[0] + ".txt"
            with open(output_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(content)

        else:
            messagebox.showerror("Erro", "Formato de conversão inválido.")
            return

        messagebox.showinfo("Sucesso", f"Arquivo convertido com sucesso: {output_path}")
    except Exception as e:
        messagebox.showerror("Erro", f"Falha ao converter o arquivo.\n{e}")


# Interface gráfica com Tkinter
root = Tk()
root.title("Conversor de Arquivos")
root.geometry("500x300")
root.config(bg="#f0f8ff")

# Variáveis
input_file = StringVar()
input_format = StringVar(value="TXT")
output_format = StringVar(value="TXT")

# Labels e Widgets
Label(root, text="Formato de Entrada:", bg="#f0f8ff", font=("Arial", 12)).pack(pady=10)
OptionMenu(root, input_format, "TXT", "DOCX").pack()

Label(root, text="Formato de Saída:", bg="#f0f8ff", font=("Arial", 12)).pack(pady=10)
OptionMenu(root, output_format, "TXT", "DOCX").pack()

Button(root, text="Selecionar Arquivo", command=select_file, bg="#add8e6", font=("Arial", 12)).pack(pady=10)
lbl_selected_file = Label(root, text="Nenhum arquivo selecionado.", bg="#f0f8ff", font=("Arial", 10))
lbl_selected_file.pack()

Button(root, text="Converter Arquivo", command=convert_file, bg="#32cd32", font=("Arial", 12)).pack(pady=20)

# Rodar a aplicação
root.mainloop()
